"""
Utility functions for file operations
"""
import os
from pathlib import Path
from typing import Dict, Any
import json
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch

def ensure_directory(path: Path) -> None:
    """Ensure a directory exists, create if it doesn't"""
    path.mkdir(parents=True, exist_ok=True)

def save_json(data: Dict[str, Any], filepath: Path) -> None:
    """Save data to JSON file"""
    ensure_directory(filepath.parent)
    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)

def load_json(filepath: Path) -> Dict[str, Any]:
    """Load data from JSON file"""
    if not filepath.exists():
        return {}
    with open(filepath, 'r', encoding='utf-8') as f:
        return json.load(f)

def docx_to_text(filepath: Path) -> str:
    """Extract text from DOCX file"""
    if not filepath.exists():
        return ""
    
    doc = Document(filepath)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text.strip())
    return '\n'.join(text)

def create_docx_report(content: str, filepath: Path, title: str = "Report") -> None:
    """Create a DOCX report with formatted content"""
    ensure_directory(filepath.parent)
    
    doc = Document()
    
    # Add title
    title_paragraph = doc.add_heading(title, 0)
    title_paragraph.alignment = 1  # Center alignment
    
    # Add content
    doc.add_paragraph(content)
    
    # Save the document
    doc.save(filepath)

def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    """Convert DOCX to PDF"""
    ensure_directory(pdf_path.parent)
    
    # Extract text from DOCX
    text = docx_to_text(docx_path)
    
    # Create PDF
    doc = SimpleDocTemplate(str(pdf_path), pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        spaceAfter=30,
        alignment=1  # Center
    )
    story.append(Paragraph("Project Report", title_style))
    story.append(Spacer(1, 12))
    
    # Add content
    content_style = ParagraphStyle(
        'Content',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6
    )
    
    # Split content into paragraphs
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip():
            story.append(Paragraph(para, content_style))
            story.append(Spacer(1, 6))
    
    doc.build(story)

def get_file_size(filepath: Path) -> str:
    """Get human-readable file size"""
    if not filepath.exists():
        return "0 B"
    
    size = filepath.stat().st_size
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size < 1024.0:
            return f"{size:.1f} {unit}"
        size /= 1024.0
    return f"{size:.1f} TB"

def cleanup_old_files(directory: Path, max_age_days: int = 7) -> None:
    """Clean up old files in a directory"""
    import time
    current_time = time.time()
    max_age_seconds = max_age_days * 24 * 60 * 60
    
    for filepath in directory.rglob("*"):
        if filepath.is_file():
            file_age = current_time - filepath.stat().st_mtime
            if file_age > max_age_seconds:
                filepath.unlink() 